import io
import pypdf
import docx

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from a PDF file using pypdf."""
    text_parts = []
    try:
        pdf_file = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error parsing PDF file: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from a DOCX file using python-docx."""
    try:
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error parsing DOCX file: {str(e)}")

def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extracts text from a TXT file using UTF-8 decoding, falling back to Latin-1."""
    try:
        return file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            return file_bytes.decode('latin-1')
        except Exception as e:
            raise ValueError(f"Error decoding TXT file: {str(e)}")

def extract_text(filename: str, file_bytes: bytes) -> str:
    """Detects file type and extracts text from the resume file bytes."""
    ext = filename.split('.')[-1].lower() if '.' in filename else ''
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        return extract_text_from_docx(file_bytes)
    elif ext in ['txt', 'md']:
        return extract_text_from_txt(file_bytes)
    else:
        # Try text decoding as general fallback
        try:
            return extract_text_from_txt(file_bytes)
        except Exception:
            raise ValueError(f"Unsupported file format: .{ext}. Supported formats are PDF, DOCX, and TXT.")
